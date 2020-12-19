from docx import Document
import os

def get_name(index):
    if index == 0:
        return '顺序高斯消去法'
    elif index == 1:
        return '选列主元高斯消去法'
    elif index == 2:
        return 'LU分解'
    elif index == 3:
        return '紧凑格式的LU分解'
    elif index == 4:
        return '二分法'
    elif index == 5:
        return '牛顿迭代法'
    elif index == 6:
        return '高斯赛德尔迭代'
    elif index == 7:
        return '拉格朗日插值'
    elif index == 8:
        return '三次样条插值'
    elif index == 9:
        return '欧拉法解常微分方程'
    elif index == 10:
        return '改进欧拉法解常微分方程'
    elif index == 11:
        return '其他'


def auto_word(index):
    savePath = './matlabcode'  # word的保存路径
    experimentName = get_name(index)  # 最大的标题
    str = ''  # 保存从txt里读出来的数据

    Doc = Document()
    Doc.styles['Normal'].font.name = u'宋体'
    Doc.add_heading("实验   " + experimentName, 0)
    Doc.add_heading('学生信息', level=2)
    stu_mess = Doc.add_paragraph()
    try:
        with open('stu_message.txt', 'r') as f:  # 读
            data = f.readlines()
            cnt = 0
            tb = ['姓名', '学号', '班级']
            for i in data:
                run = stu_mess.add_run(tb[cnt])
                cnt += 1
                run = stu_mess.add_run()
                if i[-1] == '\n':
                    run.text = i[:-1]  # 不输出最后一个字符，因为最后一个字符是换行
                else:
                    run.text = i
                run.font.underline = True
                run = stu_mess.add_run("\t")
    except:
        pass

    Doc.add_heading('一、实验目的', level=1)
    try:
        with open('mudi.txt', 'r') as f:  # 读
            data = f.read()
            str1 = data

        Doc.add_paragraph(str1)  # 写入实验目的
    except:
        pass

    Doc.add_heading('二、实验内容', level=1)
    Doc.add_paragraph('程序实现', style='List Number')
    try:
        with open('codes.txt', 'r') as f:  # 读
            data = f.read()
            str1 = data

        Doc.add_paragraph(str1)  # 写入程序实现
    except:
        pass

    Doc.add_paragraph('代码测试', style='List Number')
    try:
        with open('answer.txt', 'r') as f:  # 读
            data = f.read()
            st1r = data
        Doc.add_paragraph(st1r)  # 写入代码测试
    except:
        pass

    Doc.add_heading('三、总结和感悟', level=1)
    try:
        with open('ganwu.txt', 'r') as f:  # 读
            data = f.read()
            str1 = data
        Doc.add_paragraph(str1)  # 写入程序实现
    except:
        pass

    Doc.save("Matlab_Experimental_report.doc")


def init():
    f = open('codes1.txt', 'w')
    f.close()
    f = open('codes.txt', 'w')
    f.close()
    f = open('answer.txt', 'w')
    f.close()
    f = open('ganwu.txt', 'w')
    f.close()
    f = open('mudi.txt', 'w')
    f.close()
    f = open('stu_message.txt', 'w')
    f.close()


def bianyi(name):
    os.rename(str(name) + '.txt', str(name) + '.m')

def get_mname(index):
    path = './实验代码/'
    if index == 0:
        return path + 'nagauss.m'
    elif index == 1:
        return path + 'nagauss2.m'
    elif index == 2:
        return path + 'nalu.m'
    elif index == 3:
        return path + 'nalupad.m'
    elif index == 4:
        return path + 'nabisect.m'
    elif index == 5:
        return path + 'nanewton.m'
    elif index == 6:
        return path + 'nags.m'
    elif index == 7:
        return path + 'nalagr.m'
    elif index == 8:
        return path + 'naspline.m'
    elif index == 9:
        return path + 'naeuler.m'
    elif index == 10:
        return path + 'naeuler2.m'